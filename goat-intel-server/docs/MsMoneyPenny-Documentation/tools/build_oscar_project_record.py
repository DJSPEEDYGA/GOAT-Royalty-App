"""Build the formatted Oscar project record DOCX for verified PDF export."""

from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TABLE_HELPERS = Path(
    "/Users/raspy/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.521.10419/skills/documents/scripts"
)
sys.path.insert(0, str(TABLE_HELPERS))
from table_geometry import apply_table_geometry  # noqa: E402


ROOT = Path("/Volumes/FKD1/USB-Uncensored-LLM-main")
OUTPUT = ROOT / "Documentation" / "outputs" / "Oscar_Project_Record_and_Final_Deployment_Blueprint.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "153B60"
TEAL = "148A8A"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F6F8FA"
GREEN_FILL = "E1F1E4"
AMBER_FILL = "FFF0CC"
RED_FILL = "FCE1DB"
GRID = "C8D5DF"
TABLE_WIDTH = 9360
MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_text(cell, text, *, bold=False, color=None, size=8.2):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(document, headers, rows, widths, *, status_column=None, size=8.0):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    for index, header in enumerate(headers):
        set_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF", size=8.1)
        shade(table.rows[0].cells[index], TEAL)
        table.rows[0].cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    repeat_table_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_text(cells[index], value, size=size)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if status_column == index:
                value_lower = str(value).lower()
                if "working" in value_lower or "prepared" in value_lower:
                    shade(cells[index], GREEN_FILL)
                elif "planned" in value_lower or "to build" in value_lower:
                    shade(cells[index], AMBER_FILL)
                elif "confirm" in value_lower or "not installed" in value_lower or "waiting" in value_lower:
                    shade(cells[index], RED_FILL)
        if status_column is None and len(table.rows) % 2 == 0:
            for cell in cells:
                shade(cell, LIGHT_FILL)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=TABLE_WIDTH,
        indent_dxa=120,
        cell_margins_dxa=MARGINS,
    )
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_body(document, text, *, bold=False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def add_bullets(document, values):
    for value in values:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.12
        paragraph.add_run(value)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def prepare_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("23384A")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_document():
    document = Document()
    prepare_styles(document)
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "OSCAR PROJECT RECORD  |  FINAL DEPLOYMENT BLUEPRINT"
    header.style = document.styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(NAVY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Prepared May 25, 2026  |  Page ").font.size = Pt(8)
    add_page_field(footer)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(32)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Oscar Project Record")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    sub_run = subtitle.add_run("Final Deployment Blueprint and Source-Code Delivery Record")
    sub_run.font.size = Pt(15)
    sub_run.font.color.rgb = RGBColor.from_string(BLUE)
    add_body(document, "Prepared for the owner | Version 2026-05-25", bold=True)

    callout = document.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_text(
        callout.cell(0, 0),
        "Status: Original Oscar is local and working from the FKD1 USB source. "
        "A 52 GB NVIDIA Thor master transfer package is already built on 1TB JUMP. "
        "Final all-platform one-click delivery remains an integration and hardware-confirmation task.",
        size=10,
    )
    shade(callout.cell(0, 0), HEADER_FILL)
    apply_table_geometry(callout, [TABLE_WIDTH], table_width_dxa=TABLE_WIDTH, indent_dxa=120, cell_margins_dxa=MARGINS)
    set_cell_border(callout.cell(0, 0))

    document.add_heading("Purpose And Current Position", level=1)
    add_body(
        document,
        "This record captures what has been implemented in Oscar, what is present in the transfer "
        "master, where the editable source code lives, the requested final hardware estate, and the "
        "build requirements for a finished one-click Oscar system."
    )
    add_bullets(
        document,
        [
            "Original source: /Volumes/FKD1/USB-Uncensored-LLM-main",
            "Thor transfer master: /Volumes/1TB JUMP/Oscar-Thor-Master-USB-v1",
            "Preserved in the transfer master: UI, settings, chat history, drafts, active model store and rebuildable model files.",
            "Seven active model manifests are included, with moondream:1.8b available for image understanding.",
            "Graphics generation, 3D production, broad audio/video intake and distributed orchestration are future implementation work.",
        ],
    )

    document.add_heading("Completed Work", level=1)
    add_table(
        document,
        ["Date", "Work Completed", "Verified Outcome", "State"],
        [
            ["2026-05-22 to 05-23", "Persistent memory, identity guidance and Tool Mode foundations", "Settings and workflow stored with USB app", "Working"],
            ["2026-05-23", "Mac launcher and portable-client packaging", "Launcher guidance and portable behavior recorded", "Working"],
            ["2026-05-25", "Voice conversation upgrade", "Dictation, read-aloud, wake listening and saved profile", "Working"],
            ["2026-05-25", "Skill/model routing and vision addition", "Chat/Voice/Code/Research/Vision modes; moondream manifest", "Working"],
            ["2026-05-25", "Crew feature merge and mobile repair", "Expert views, Council review and phone layout", "Working"],
            ["2026-05-25", "NVIDIA Thor master transfer build", "52 GB package copied and key files checked", "Prepared"],
        ],
        [1600, 2700, 3460, 1600],
        status_column=3,
        size=7.8,
    )

    document.add_heading("Capability Truth Table", level=1)
    add_table(
        document,
        ["Capability", "Current State", "Evidence", "Expansion Needed"],
        [
            ["Chat and saved memory", "Working", "UI, local server and saved settings", "Searchable retrieval layer"],
            ["Voice conversation", "Working through browser voice", "UI and voice upgrade draft", "Licensed/consented neural voice"],
            ["Image understanding", "Working, host-dependent", "moondream:1.8b manifest", "Higher-performance visual models"],
            ["PDF intake", "Working with text snippet", "Roughly first 12,000 extracted characters", "Indexed document vault and citations"],
            ["Code/research roles", "Working as local workflow", "Skills, Crew panel and Tool Mode", "Larger models and adapters"],
            ["Image/graphic generation", "Not installed", "No engine wired into Oscar", "CUDA generation workflow"],
            ["3D art/rendering", "Not installed", "No Blender bridge", "Controlled 3D runner and preview"],
            ["Audio/music and video intake", "Not installed", "UI accepts images/PDF only", "Upload, transcription and FFmpeg"],
            ["Autonomous production work", "Partially designed, constrained", "Approved local actions required", "Audited queues and permissions"],
        ],
        [2050, 2050, 2710, 2550],
        status_column=1,
        size=7.6,
    )

    document.add_heading("Reported Final Hardware Estate", level=1)
    add_body(
        document,
        "Hardware statements below preserve the owner's requested targets while marking details that "
        "must be confirmed before the final installers and launchers are locked."
    )
    add_table(
        document,
        ["Node", "Reported Hardware Or Storage", "Oscar Role", "Verification Status"],
        [
            ["Oscar final PC home", "42-core AMD Threadripper; 256 GB RAM; 2 x RTX 5090 Founders Edition; 100 TB internal 10-SSD PCIe storage; 100 TB Thunderbolt external cloud rack", "Creative production and CUDA/WSL services", "Confirm CPU, SSD layout and runtime"],
            ["Jetson AGX Thor", "Purchased; official kit reference states T5000 and 128 GB LPDDR5X", "Always-on service and orchestration hub", "Thor scripts prepared"],
            ["Mac Studio", "Requested as Mac M5 128 G", "macOS creator access and launch target", "Confirm exact model/chip/RAM"],
            ["Secondary Jetson", "Requested as Jetson Nano 64 G with 2 TB", "Satellite service node", "Confirm exact NVIDIA SKU"],
            ["Server vault", "100 TB storage reported", "Media vault, projects, assets and backups", "Confirm filesystem and permissions"],
        ],
        [1820, 3070, 2260, 2210],
        status_column=3,
        size=7.55,
    )
    add_body(
        document,
        "Important confirmation point: current official Apple and NVIDIA product pages should be "
        "checked against the exact purchased Mac Studio and secondary Jetson SKU before final one-click installers are finalized."
    )

    document.add_heading("Recommended Multi-Node Architecture", level=1)
    add_table(
        document,
        ["Layer", "Primary Node", "Function"],
        [
            ["Control and orchestration", "Thor", "Always-on chat/service hub, routing, schedules, indexes and permissions"],
            ["Creative generation", "Oscar final PC home", "Artwork, audio/video processing, 3D rendering and CUDA workflows"],
            ["macOS creator access", "Mac Studio", "Review, creation workflows and macOS launch experience"],
            ["Satellite workflow", "Confirmed Jetson SKU", "Scaled edge/local services after capability verification"],
            ["Durable storage", "100 TB server", "Originals, masters, assets, reference library and backups"],
            ["Recovery and transfer", "Oscar master USB", "Source package, scripts, model seed and recovery workflow"],
        ],
        [2200, 2250, 4910],
        size=7.9,
    )

    document.add_heading("Editable Source-Code Inventory", level=1)
    add_body(
        document,
        "The final Oscar delivery includes editable source code and rebuild tools, rather than only packaged executables."
    )
    add_table(
        document,
        ["Component", "Live Source Location", "Purpose", "Final Package"],
        [
            ["Web interface", "Shared/FastChatUI.html", "Voice, skills, Crew and attachments", "Include"],
            ["Local server", "Shared/chat_server.py", "Persistence, API proxy and tools", "Include"],
            ["Settings/memory", "Shared/chat_data/settings.json", "Identity and saved choices", "Include with privacy review"],
            ["Chat history", "Shared/chat_data/chats.json", "Saved conversations", "Owner-selectable export"],
            ["Mac launchers", "Launch Oscar.command; Mac/*.command", "macOS baseline launch", "Include"],
            ["Windows scripts", "Windows/install*; Windows/start-fast-chat.bat", "Windows portable baseline", "Include"],
            ["Thor scripts", "Thor/install*, start*, verify*, build-master*", "ARM64/JetPack delivery path", "Include"],
            ["Documentation tools", "Documentation/tools/*", "Rebuild PDF and XLSX record", "Include"],
        ],
        [1750, 3450, 2760, 1400],
        size=7.45,
    )

    document.add_heading("Final One-Click Delivery Requirement", level=1)
    add_body(
        document,
        "A final platform is only ready when it has one obvious launch entry point, its editable source, "
        "prerequisite checks, diagnostics, migration behavior and acceptance tests."
    )
    add_table(
        document,
        ["Platform", "Required One-Click Result", "Current Position", "Acceptance Focus"],
        [
            ["Oscar final PC home", "Creative node launch and vault connection", "Baseline exists; CUDA/WSL launcher to build", "GPU, media jobs, storage and recovery"],
            ["Mac Studio", "Apple Silicon client/service launch", "Baseline exists; test after exact confirmation", "Runtime, voice and vault access"],
            ["Jetson AGX Thor", "ARM64 setup once; direct service launch", "Transfer scripts prepared", "JetPack, inference and restart"],
            ["Secondary Jetson", "Scaled satellite launch", "Waiting for exact SKU", "Memory fit and performance"],
            ["Server vault", "Mount and validate private data paths", "Service/permission plan to build", "Backup, indexing and privacy"],
        ],
        [1850, 2810, 2570, 2130],
        status_column=2,
        size=7.5,
    )
    document.add_heading("Every Final Package Must Include", level=2)
    add_bullets(
        document,
        [
            "Editable source files and source manifest.",
            "Prerequisite installer or system check plus one obvious launcher.",
            "Settings and memory migration path with privacy boundaries.",
            "Model/runtime placement, media-vault connection and GPU/storage diagnostics.",
            "Logs, backup/recovery directions and repeatable acceptance tests for chat, voice, tools, files and services.",
        ],
    )

    document.add_heading("Media, Art, Music And 3D Roadmap", level=1)
    add_table(
        document,
        ["Phase", "Capability", "Needed Runtime Or Tool Class", "Target Node", "State"],
        [
            ["1", "Media Vault and indexing", "Uploads, metadata DB and retrieval/index", "Thor + server", "Planned"],
            ["2", "Audio/video intake", "Resumable upload, FFmpeg and transcription", "Windows + vault", "Planned"],
            ["3", "Graphics studio", "CUDA image generation, edit/upscale/export", "Oscar final PC home", "Planned"],
            ["4", "Music workspace", "Catalog, stems, notes and rights tracking", "Windows + vault", "Planned"],
            ["5", "3D studio", "Blender runner, previews and rendering", "Windows", "Planned"],
            ["6", "Larger model routing", "Benchmarked code, vision and multimodal models", "Thor + Windows", "Planned"],
            ["7", "Audited automation", "Owner-approved jobs, checkpoints and logs", "Thor", "Planned"],
        ],
        [670, 2080, 3680, 1690, 1240],
        status_column=4,
        size=7.6,
    )

    document.add_heading("Can Oscar Produce Work Like This?", level=1)
    add_body(
        document,
        "Today Oscar has chat, saved memory, browser voice controls, image understanding, PDF snippet "
        "intake and constrained local-tool workflow. He does not yet independently expose a report, "
        "PDF or spreadsheet-generation tool."
    )
    add_body(
        document,
        "On the final system, yes: Oscar can be equipped to create structured records, PDF exports, "
        "spreadsheets, asset inventories and project packages by adding audited document and spreadsheet "
        "adapters through Tool Mode. The same principle applies to graphics, music and 3D work: capabilities "
        "become real after the engine, permissions, storage and acceptance tests are connected and verified."
    )

    document.add_heading("Source References", level=1)
    add_table(
        document,
        ["Reference", "Use"],
        [
            ["https://www.nvidia.com/en-us/autonomous-machines/embedded-systems/jetson-thor/", "Official Jetson Thor product reference"],
            ["https://docs.nvidia.com/jetson/agx-thor-devkit/user-guide/latest/index.html", "Official Thor setup path"],
            ["https://www.nvidia.com/en-us/autonomous-machines/embedded-systems/jetson-orin", "Official Jetson family comparison"],
            ["https://www.apple.com/mac-studio/specs/", "Official current Mac Studio reference"],
            ["https://docs.ollama.com/faq", "Ollama model loading and multi-GPU reference"],
            ["https://docs.nvidia.com/cuda/wsl-user-guide/", "Windows WSL/CUDA reference"],
            ["/Volumes/FKD1/USB-Uncensored-LLM-main", "Live Original Oscar source"],
            ["/Volumes/1TB JUMP/Oscar-Thor-Master-USB-v1", "Completed transfer master artifact"],
        ],
        [6200, 3160],
        size=7.5,
    )

    document.add_heading("Record Maintenance", level=1)
    add_body(
        document,
        "Keep this record and its cross-reference workbook with the Oscar master USB. Update both after "
        "the exact Mac Studio and secondary Jetson hardware are confirmed, after each final launcher is built, "
        "and after new media, graphics or 3D capabilities pass acceptance tests."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(f"DOCX exported: {output}")
